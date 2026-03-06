#!/usr/bin/env python3
"""
corrector.py — Corrector de format per a la revista InDret (v2)
Ús:     python3 corrector.py <article.docx>
Sortida: <nom>_corregit.docx  +  <nom>_informe.md
"""

import sys, re, os, argparse, zipfile, io
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree as _lxml_et

# ─── Constants ───────────────────────────────────────────────────────────────
FONT_SERIF = "PT Serif"
FONT_SANS  = "Open Sans"

# Sagnat de l'índex per nivell
INDEX_INDENT = {1: Cm(0.0), 2: Cm(0.5), 3: Cm(1.0), 4: Cm(1.5)}

# Abreviatures legals que NO s'han de convertir a versaletes
LEGAL_ABBR = frozenset({
    'STS', 'STSJ', 'SAP', 'SJPI', 'SJP', 'STC', 'STJUE', 'TJUE',
    'TC', 'TS', 'BOE', 'CC', 'CP', 'CE', 'LEC', 'LO', 'PE', 'UE',
    'EU', 'ECLI', 'ROJ', 'BGB', 'STGB', 'BGH',
    'ET AL', 'OP CIT', 'NM', 'LH', 'FS', 'ART',
    'I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X',
})

# ─── Patrons regex ───────────────────────────────────────────────────────────
RE_DOUBLE_SPACE = re.compile(r'  +')
RE_QUOT_OPEN    = re.compile(r'(?<!«)"(\S)')
RE_QUOT_CLOSE   = re.compile(r'(\S)"(?!»)')
RE_H3    = re.compile(r'^\d+\.\d+\.\d+\.?\s+\S')   # 1.1.1.
RE_H2    = re.compile(r'^\d+\.\d+\.?\s+\S')         # 1.1.
RE_H1    = re.compile(r'^\d+\.\s+\S')               # 1.
RE_H4    = re.compile(r'^[a-z]\.\s+\S')             # a.
RE_JURIS  = re.compile(r'\b(STS|STSJ|SAP|SJPI|SJP|STC|STJUE)\b')
RE_ECLI   = re.compile(r'ECLI:[A-Z]{2}:[^:]+:\d+:[A-Z0-9.]+')
RE_ROJ    = re.compile(r'Roj\s*:\s*\S+')
RE_OP_CIT = re.compile(r'\bop\.?\s*cit\.?\b', re.IGNORECASE)
RE_CAPS   = re.compile(r'^[A-ZÁÉÍÓÚÑÜÀÈÌÒÙ\s\-]+$')

BIB_KEYWORDS   = {'bibliografía', 'referencias bibliográficas', 'bibliography',
                  'referencias', 'bibliografia'}
INDEX_KEYWORDS = {'índice', 'sumario', 'tabla de contenido', 'sumari', 'índex'}

# Títols N1 que NO s'han de numerar (bibliografía i equivalents)
NO_NUMBER_HEADINGS = frozenset({
    'bibliografía', 'bibliografía y referencias', 'referencias bibliográficas',
    'referencias', 'bibliography', 'referencias y bibliografía',
    'bibliografia', 'bibliografía general',
})


# ─── Helpers XML/tipogràfics ─────────────────────────────────────────────────

def get_heading_level(text: str) -> int:
    """Retorna el nivell del títol (1-4) o 0 si no és títol."""
    if RE_H3.match(text): return 3
    if RE_H2.match(text): return 2
    if RE_H1.match(text): return 1
    if RE_H4.match(text): return 4
    return 0


def classify_para(para) -> str:
    text = para.text.strip()
    if not text: return 'empty'
    tl = text.lower()
    if tl in BIB_KEYWORDS:   return 'bib_keyword'
    if tl in INDEX_KEYWORDS: return 'index_keyword'
    level = get_heading_level(text)
    if level: return f'h{level}'
    # Estil Word explícit
    sn = (para.style.name or '').lower()
    for lv, variants in [
        ('h1', ['heading 1', 'título 1', 'encabezado 1', 'títol 1']),
        ('h2', ['heading 2', 'título 2', 'encabezado 2']),
        ('h3', ['heading 3', 'título 3', 'encabezado 3']),
    ]:
        if any(v in sn for v in variants):
            return lv
    # Fallback: paràgraf curt i tot en negreta → N1 (articles sense numeració)
    if text and len(text) < 120 and para.runs:
        bold_runs = [r for r in para.runs if r.text.strip()]
        if bold_runs and all(bool(r.bold) for r in bold_runs):
            return 'h1'
    return 'body'


def is_likely_surname(text: str) -> bool:
    """True si el text és un cognom en MAJÚSCULES que cal convertir a versaletes."""
    t = text.strip().rstrip('.,;:')
    if not t or len(t) < 3:
        return False
    if not RE_CAPS.match(t):          # ha de ser tot majúscules/espais/guions
        return False
    if t != t.upper():                 # ja ha de ser majúscules
        return False
    key = re.sub(r'[\s.\-]', '', t.upper())
    if t.upper().strip() in LEGAL_ABBR or key in {
        re.sub(r'[\s.\-]', '', a) for a in LEGAL_ABBR
    }:
        return False
    if ' ' not in t and len(t) < 4:   # paraula sola curta → descarta
        return False
    return True


def set_run_font(run, font_name: str, size_pt, bold=None, italic=None):
    """Aplica font i mida. Si bold/italic és None, no els toca."""
    run.font.name = font_name
    run.font.size = size_pt
    if bold   is not None: run.font.bold   = bold
    if italic is not None: run.font.italic = italic
    rpr = run._r.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.insert(0, rf)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rf.set(qn(attr), font_name)
    # Elimina atributs de font temàtica que Word prioritza sobre w:ascii/w:hAnsi
    for theme_attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme', 'w:eastAsiaTheme'):
        if rf.get(qn(theme_attr)) is not None:
            del rf.attrib[qn(theme_attr)]


def set_small_caps_xml(r_elem):
    """Aplica versaletes a un element XML <w:r>."""
    rpr = r_elem.find(qn('w:rPr'))
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        r_elem.insert(0, rpr)
    for e in rpr.findall(qn('w:smallCaps')):
        rpr.remove(e)
    sc = OxmlElement('w:smallCaps')
    sc.set(qn('w:val'), '1')
    rpr.append(sc)


def set_line_spacing(para, value: float):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = value


def add_bookmark(para, bm_id: int, bm_name: str):
    """Afegeix bookmark XML al paràgraf (al voltant dels runs)."""
    p = para._p
    runs = p.findall(qn('w:r'))
    if not runs:
        return
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'),   str(bm_id))
    bm_start.set(qn('w:name'), bm_name)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bm_id))
    children = list(p)
    idx_first = children.index(runs[0])
    idx_last  = children.index(runs[-1])
    p.insert(idx_first, bm_start)
    p.insert(idx_last + 2, bm_end)   # +2 per l'element que acabem d'inserir


def add_internal_hyperlink(para, anchor: str):
    """Embolcalla els runs del paràgraf en un <w:hyperlink> intern."""
    p = para._p
    runs = p.findall(qn('w:r'))
    if not runs:
        return
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), anchor)
    children = list(p)
    insert_idx = children.index(runs[0])
    for run in runs:
        p.remove(run)
        # Afegir estil d'hiperenllaç (blau subratllat)
        rpr = run.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            run.insert(0, rpr)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rpr.append(color)
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rpr.append(u)
        hl.append(run)
    p.insert(insert_idx, hl)


# ─── ZIP helpers per injectar footnotes al document de sortida ───────────────

def _ensure_footnotes_rel(rels_data: bytes) -> bytes:
    """Afegeix la relació footnotes a document.xml.rels si no hi existeix."""
    text = rels_data.decode('utf-8')
    if 'relationships/footnotes' in text:
        return rels_data
    ids = re.findall(r'Id="rId(\d+)"', text)
    nxt = max((int(i) for i in ids), default=0) + 1
    rel = (
        f'<Relationship Id="rId{nxt}" '
        f'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" '
        f'Target="footnotes.xml"/>'
    )
    return text.replace('</Relationships>', f'  {rel}\n</Relationships>').encode('utf-8')


def _ensure_footnotes_content_type(ct_data: bytes) -> bytes:
    """Afegeix el tipus de contingut per footnotes a [Content_Types].xml si no hi és."""
    text = ct_data.decode('utf-8')
    if 'footnotes' in text.lower():
        return ct_data
    override = (
        '<Override PartName="/word/footnotes.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
    )
    return text.replace('</Types>', f'  {override}\n</Types>').encode('utf-8')


def _merge_footnotes_xml(tmpl_fn_data: bytes, article_fn_data: bytes) -> bytes:
    """Substitueix les notes de la plantilla (id≥1) per les de l'article."""
    W_FN = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnote'
    W_ID = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id'
    tmpl_root    = _lxml_et.fromstring(tmpl_fn_data)
    article_root = _lxml_et.fromstring(article_fn_data)
    # Eliminar notes demo de la plantilla (id≥1); conservar separadors (id≤0)
    for fn in list(tmpl_root):
        if fn.tag == W_FN:
            try:
                if int(fn.get(W_ID, 0)) >= 1:
                    tmpl_root.remove(fn)
            except (ValueError, TypeError):
                pass
    # Afegir notes de l'article (id≥1)
    for fn in article_root:
        if fn.tag != W_FN:
            continue
        try:
            if int(fn.get(W_ID, 0)) >= 1:
                tmpl_root.append(deepcopy(fn))
        except (ValueError, TypeError):
            pass
    return _lxml_et.tostring(tmpl_root, xml_declaration=True,
                              encoding='UTF-8', standalone=True)


def _fix_footnote_style_spacing(styles_data: bytes) -> bytes:
    """Elimina space_before/after de l'estil FootnoteText a styles.xml."""
    root = _lxml_et.fromstring(styles_data)
    W_STYLE   = qn('w:style')
    W_NAME    = qn('w:name')
    W_TYPE    = qn('w:type')
    W_PPR     = qn('w:pPr')
    W_SPACING = qn('w:spacing')
    W_VAL     = qn('w:val')
    # Noms possibles de l'estil de nota al peu (en, es, ca…)
    FOOTNOTE_NAMES = {'footnote text', 'nota al pie', 'notaalpie', 'nota al peu',
                      'footnotetext', 'footnotestyle'}
    for style in root.iter(W_STYLE):
        if style.get(W_TYPE) != 'paragraph':
            continue
        name_el = style.find(W_NAME)
        if name_el is None:
            continue
        name_val = name_el.get(W_VAL, '').lower().replace('\xa0', ' ').strip()
        if name_val not in FOOTNOTE_NAMES:
            continue
        pPr = style.find(W_PPR)
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            style.append(pPr)
        sp = pPr.find(W_SPACING)
        if sp is None:
            sp = OxmlElement('w:spacing')
            pPr.append(sp)
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'), '40')   # 2pt entre notes
        sp.set(qn('w:line'), '212')   # interlineat ~1.0 per 8.5pt
        sp.set(qn('w:lineRule'), 'auto')
    return _lxml_et.tostring(root, xml_declaration=True,
                             encoding='UTF-8', standalone=True)


def _inject_footnotes(output_path: str, article_fn_xml: bytes,
                      fn_rels: bytes | None = None):
    """Copia les notes al peu de l'article al document de sortida (post-save ZIP)."""
    RELS_PATH = 'word/_rels/footnotes.xml.rels'
    buf = io.BytesIO()
    with zipfile.ZipFile(output_path, 'r') as zin:
        has_fn   = 'word/footnotes.xml' in zin.namelist()
        has_rels = RELS_PATH in zin.namelist()
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/footnotes.xml':
                    data = _merge_footnotes_xml(data, article_fn_xml)
                elif item.filename == 'word/styles.xml':
                    data = _fix_footnote_style_spacing(data)
                elif item.filename == 'word/_rels/document.xml.rels':
                    data = _ensure_footnotes_rel(data)
                elif item.filename == '[Content_Types].xml':
                    data = _ensure_footnotes_content_type(data)
                elif item.filename == RELS_PATH and fn_rels is not None:
                    data = fn_rels  # substituïm les rels de la plantilla
                zout.writestr(item, data)
            if not has_fn:
                zout.writestr('word/footnotes.xml', article_fn_xml)
            if not has_rels and fn_rels is not None:
                zout.writestr(RELS_PATH, fn_rels)
    with open(output_path, 'wb') as f:
        f.write(buf.getvalue())


_RE_IDX_PREFIX = re.compile(r'^(\d+(?:\.\d+)*)[.\t\s]')


def _index_level_from_text(text: str) -> int:
    """Detecta el nivell d'una entrada d'índex pel prefix numèric del text.

    "1. Introducción" → 1
    "2.1. Costes"     → 2
    "2.1.1. Riesgo"   → 3
    "Bibliografía"    → 1 (sense prefix → nivell 1)
    """
    m = _RE_IDX_PREFIX.match(text)
    if not m:
        return 1
    return len(m.group(1).split('.'))


_INDEX_TEXT_START = 720  # Twips: columna on comença el text de TOTS els nivells


def _build_index_para_elem(text: str, level: int):
    """Crea un element <w:p> net per a una entrada d'índex.

    Tots els nivells comparteixen el mateix text_start (hanging indent):
      - El número surt des del marge (first line indent = 0)
      - El text comença a _INDEX_TEXT_START
      - Les segones línies tornen a _INDEX_TEXT_START (wrap correcte)
    Nivell 1: negreta cursiva. Nivell 2+: normal.
    """
    # Normalitzar separador: "2.1. Texto" → "2.1.\tTexto" per al tab stop
    text_norm = re.sub(r'^(\d[\d.]*)\.\s+', r'\1.\t', text)

    p = OxmlElement('w:p')

    pPr = OxmlElement('w:pPr')

    sp = OxmlElement('w:spacing')
    sp.set(qn('w:line'), '240')
    sp.set(qn('w:lineRule'), 'auto')
    sp.set(qn('w:after'), '0')
    sp.set(qn('w:before'), '0')
    pPr.append(sp)

    # Hanging indent: número al marge, text i wraps a _INDEX_TEXT_START
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'),    str(_INDEX_TEXT_START))
    ind.set(qn('w:hanging'), str(_INDEX_TEXT_START))
    pPr.append(ind)

    # Tab stop únic per a tots els nivells
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'left')
    tab.set(qn('w:pos'), str(_INDEX_TEXT_START))
    tabs.append(tab)
    pPr.append(tabs)

    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_SERIF)
    rFonts.set(qn('w:hAnsi'), FONT_SERIF)
    rFonts.set(qn('w:cs'), FONT_SERIF)
    rPr.append(rFonts)
    if level == 1:
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:i'))
    for tag in ('w:sz', 'w:szCs'):
        el = OxmlElement(tag)
        el.set(qn('w:val'), '18')   # 9pt
        rPr.append(el)
    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = text_norm
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p


# ─── Extractor de metadades ──────────────────────────────────────────────────
class MetadataExtractor:
    """Extreu metadades i seccions d'un article per omplir la plantilla."""

    RE_SUMARI    = re.compile(r'^(RESUMEN|SUMARIO|SUMARI|RESUM)\s*:\s*(.*)', re.I | re.S)
    RE_ABSTRACT  = re.compile(r'^ABSTRACT\s*:\s*(.*)', re.I | re.S)
    RE_KW_ES     = re.compile(r'^(Palabras?\s+clave|Paraules\s+clau|Palabras-clave)\s*:\s*(.*)', re.I)
    RE_KW_EN     = re.compile(r'^Keywords?\s*:\s*(.*)', re.I)
    RE_INDEX_HDR = re.compile(r'^(ÍNDICE|ÍNDEX|SUMARIO|SUMARI|TABLA\s+DE\s+CONTENIDO)\s*$', re.I)
    RE_BIB_HDR   = re.compile(r'^(Bibliografía|Referencias\s+bibliográficas|Bibliography|Referencias|Bibliografia)\s*$', re.I)

    # Regex per eliminar prefix numèric d'un títol ("2.1.3. Texto" → "Texto")
    RE_NUM_PREFIX = re.compile(r'^\d+(\.\d+)*\.?\s+')

    @staticmethod
    def _detect_index_level(para) -> int:
        """Detecta el nivell d'una entrada d'índex pel format del paràgraf."""
        text_runs = [r for r in para.runs if r.text.strip()]
        if not text_runs:
            return 1
        if any(r.bold for r in text_runs):
            return 1   # N1: negreta
        if any(r.italic for r in text_runs):
            return 2   # N2: qualsevol run en cursiva
        return 3       # sense format → N3

    def _refine_index_levels(self, index_entries: list, paras: list,
                              body_start: int, body_end: int) -> list:
        """Millora els nivells de l'índex buscant els títols al cos de l'article.
        El cos té numeració (2.1., 2.1.1.) que permet saber el nivell real."""
        # Construïm mapa: text_normalitzat → nivell (del cos)
        level_map = {}
        for p in paras[body_start:body_end]:
            text = p.text.strip()
            if not text:
                continue
            lvl = get_heading_level(text)
            if lvl and lvl <= 4:
                normalized = self.RE_NUM_PREFIX.sub('', text)
                level_map[normalized] = lvl

        if not level_map:
            return index_entries

        refined = []
        for text, lvl in index_entries:
            body_lvl = level_map.get(text)
            refined.append((text, body_lvl if body_lvl else lvl))
        return refined

    def _is_metadata_line(self, text: str) -> bool:
        return bool(
            self.RE_SUMARI.match(text) or self.RE_ABSTRACT.match(text) or
            self.RE_KW_ES.match(text)  or self.RE_KW_EN.match(text) or
            self.RE_INDEX_HDR.match(text)
        )

    # Capçaleres de secció que indiquen portada en taula (article ja formatat)
    _PORTADA_SECTION_HDRS = re.compile(
        r'^(Sumario|Sumari|Resumen|Resum|Abstract)$', re.I
    )
    _PORTADA_KW_HDRS = re.compile(
        r'^(Title|Título en inglés|Palabras?\s+clave|Paraules\s+clau|Keywords?)\s*:',
        re.I
    )

    def _try_portada_tables(self, doc) -> dict | None:
        """Intenta extreure metadades de la taula de portada (article ja formatat).

        Retorna un dict parcial si s'ha detectat estructura de portada InDret
        (Taula 0 amb Sumario/Abstract, Taula 1 amb Índex), o None si no s'ha
        detectat.
        """
        if not doc.tables:
            return None
        t0 = doc.tables[0]
        if not t0.rows or len(t0.rows[0].cells) < 2:
            return None

        cell_right = t0.rows[0].cells[1]
        cell_left  = t0.rows[0].cells[0] if len(t0.rows[0].cells) > 0 else None

        # Detectar si la cel·la dreta té estructura de portada InDret
        cell_texts = [p.text.strip() for p in cell_right.paragraphs]
        has_portada = any(self._PORTADA_SECTION_HDRS.match(t) for t in cell_texts if t)
        if not has_portada:
            return None

        result: dict = {
            'titol': '', 'subtitol': '', 'autor': '', 'organitzacio': '',
            'sumari': [], 'abstract': [], 'titol_en': '',
            'paraules_clau': '', 'keywords': '',
            'index_entries': [], 'index_paras': [],
            'body_start_idx': 0, 'body_end_idx': len(doc.paragraphs),
        }

        # Extreure títol, sumari, abstract, keywords de la cel·la dreta
        state = 'pre_title'
        for p in cell_right.paragraphs:
            text = p.text.strip()
            if not text or text == '-':
                continue
            tl = text.lower()
            if state == 'pre_title' and not self._PORTADA_SECTION_HDRS.match(text):
                if not result['titol']:
                    result['titol'] = text
                continue
            if self._PORTADA_SECTION_HDRS.match(text):
                if any(k in tl for k in ('sumario', 'sumari', 'resumen', 'resum')):
                    state = 'sumari'
                elif 'abstract' in tl:
                    state = 'abstract'
                continue
            if self._PORTADA_KW_HDRS.match(text):
                m = re.match(r'^(Title|Título en inglés)\s*:\s*(.*)', text, re.I)
                if m:
                    result['titol_en'] = m.group(2).strip()
                    state = 'kw'
                    continue
                m = re.match(r'^(Palabras?\s+clave|Paraules\s+clau)\s*:\s*(.*)', text, re.I)
                if m:
                    result['paraules_clau'] = m.group(2).strip()
                    state = 'kw'
                    continue
                m = re.match(r'^Keywords?\s*:\s*(.*)', text, re.I)
                if m:
                    result['keywords'] = m.group(1).strip()
                    state = 'kw'
                    continue
                state = 'kw'
                continue
            if state == 'sumari':
                result['sumari'].append(text)
            elif state == 'abstract':
                result['abstract'].append(text)

        # Extreure autors de la cel·la esquerra
        if cell_left:
            lines = [p.text.strip() for p in cell_left.paragraphs
                     if p.text.strip() and p.text.strip() not in ('-',) and
                     not re.match(r'^\d+\.\d+$', p.text.strip())]
            if lines:
                result['autor'] = lines[0]
            if len(lines) > 1:
                result['organitzacio'] = lines[1]

        # Extreure índex de la taula 1 (si existeix)
        if len(doc.tables) > 1:
            t1 = doc.tables[1]
            if t1.rows and len(t1.rows[0].cells) >= 2:
                idx_cell = t1.rows[0].cells[1]
                in_index = False
                for p in idx_cell.paragraphs:
                    text = p.text.strip()
                    if not text or text == '-':
                        continue
                    if re.match(r'^(Índice|Índex|Sumario|Sumari)$', text, re.I):
                        in_index = True
                        continue
                    if in_index:
                        lvl = _index_level_from_text(text)
                        result['index_entries'].append((text, lvl))
                        result['index_paras'].append(
                            _build_index_para_elem(text, lvl)
                        )

        return result if (result['titol'] or result['sumari']) else None

    def extract(self, doc) -> dict:
        paras = doc.paragraphs
        n     = len(paras)
        data  = {
            'titol': '', 'subtitol': '', 'autor': '', 'organitzacio': '',
            'sumari': '', 'abstract': '', 'titol_en': '',
            'paraules_clau': '', 'keywords': '',
            'index_entries': [],
            'index_paras':   [],
            'body_start_idx': 0,
            'body_end_idx': n,
        }

        # Intentar extracció de la taula de portada (article ja formatat)
        portada = self._try_portada_tables(doc)
        if portada is not None:
            return portada

        # Pas 1: títol = primer paràgraf no buit
        title_idx = -1
        for i, p in enumerate(paras):
            if p.text.strip():
                data['titol'] = p.text.strip()
                title_idx = i
                break
        if title_idx < 0:
            return data

        # Pas 2: màquina d'estats per extreure seccions
        state           = 'post_title'
        sumari_parts    = []
        abstract_parts  = []
        index_entries   = []
        index_start_idx = -1
        seen_headings   = {}   # text → primer índex

        i = title_idx + 1
        while i < n:
            p    = paras[i]
            text = p.text.strip()
            if not text:
                i += 1
                continue

            m_sum = self.RE_SUMARI.match(text)
            m_abs = self.RE_ABSTRACT.match(text)
            m_kes = self.RE_KW_ES.match(text)
            m_ken = self.RE_KW_EN.match(text)
            m_idx = self.RE_INDEX_HDR.match(text)

            if m_sum:
                state = 'sumari'
                rest  = m_sum.group(2).strip()
                if rest:
                    sumari_parts.append(rest)
            elif m_abs:
                state = 'abstract'
                # Paràgraf just abans de l'ABSTRACT → candidat a títol anglès
                for j in range(i - 1, title_idx, -1):
                    prev = paras[j].text.strip()
                    if prev and not self._is_metadata_line(prev) and prev != data['titol']:
                        data['titol_en'] = prev
                        break
                rest = m_abs.group(1).strip()
                if rest:
                    abstract_parts.append(rest)
            elif m_kes:
                state = 'kw_es'
                data['paraules_clau'] = m_kes.group(2).strip()
            elif m_ken:
                state = 'kw_en'
                data['keywords'] = m_ken.group(1).strip()
            elif m_idx:
                state           = 'index'
                index_start_idx = i
            elif state == 'sumari':
                if not self._is_metadata_line(text):
                    sumari_parts.append(text)
            elif state == 'abstract':
                abstract_parts.append(text)
            elif state == 'index':
                # Detectem fi de l'índex: repetició d'un títol ja vist
                if text in seen_headings and seen_headings[text] > index_start_idx:
                    data['body_start_idx'] = i
                    state = 'body'
                else:
                    lvl = self._detect_index_level(p)
                    index_entries.append((text, lvl))
                    data['index_paras'].append(deepcopy(p._p))
            elif state == 'post_title':
                # Paràgrafs curts entre el títol i el RESUMEN → autor/org
                if len(text) < 120 and not self._is_metadata_line(text):
                    if not data['autor']:
                        data['autor'] = text
                    elif not data['organitzacio']:
                        data['organitzacio'] = text

            # Registrar primeres ocurrències per detectar inici del cos
            if text not in seen_headings:
                seen_headings[text] = i

            i += 1

        data['sumari']   = sumari_parts    # llista de paràgrafs
        data['abstract'] = abstract_parts  # llista de paràgrafs

        # Pas 3: fallback si body_start no s'ha detectat inline
        if data['body_start_idx'] == 0 and index_start_idx >= 0:
            index_set = {t for (t, _) in index_entries if t}
            for i in range(index_start_idx + 1, n):
                t = paras[i].text.strip()
                if t and t in index_set:
                    first_occ = seen_headings.get(t, i)
                    if first_occ <= index_start_idx + len(index_entries) + 5 and i > first_occ:
                        data['body_start_idx'] = i
                        break

        # Pas 4: fi del cos = última línia (inclou bibliografía)
        data['body_end_idx'] = n

        data['index_entries'] = index_entries
        return data


# ─── Numeració de pàgines ─────────────────────────────────────────────────────

def _add_page_numbers(doc, body_start_page: int = 1):
    """Afegeix numeració de pàgines al peu, centrada, Open Sans 10pt.

    Les seccions de portada i índex no mostren número (footers vinculats buits).
    La darrera secció (cos de l'article) mostra el número de pàgina,
    començant a body_start_page (default 1) i incrementant consecutivament.
    El camp PAGE s'actualitza sol quan l'usuari obre/imprimeix el document.
    Per canviar el número d'inici cal usar el paràmetre body_start_page (des del
    formulari web) — NO editar el número al peu directament.
    """
    sections = doc.sections
    if not sections:
        return

    # Darrera secció = cos de l'article: desenllaçar i afegir numeració
    body_section = sections[-1]
    body_section.footer.is_linked_to_previous = False

    # Forçar pgNumType start=body_start_page al sectPr del cos.
    # IMPORTANT: l'OOXML exigeix un ordre concret de fills a sectPr;
    # pgNumType ha d'anar ABANS de cols, titlePg, docGrid, etc.
    sect_pr = body_section._sectPr
    pg_num_type = sect_pr.find(qn('w:pgNumType'))
    if pg_num_type is None:
        pg_num_type = OxmlElement('w:pgNumType')
        # Elements que han d'anar DESPRÉS de pgNumType segons l'esquema OOXML
        _after = {qn('w:cols'), qn('w:formProt'), qn('w:vAlign'),
                  qn('w:noEndnote'), qn('w:titlePg'), qn('w:textDirection'),
                  qn('w:bidi'), qn('w:rtlGutter'), qn('w:docGrid'),
                  qn('w:printerSettings'), qn('w:sectPrChange')}
        insert_before = next((c for c in sect_pr if c.tag in _after), None)
        if insert_before is not None:
            insert_before.addprevious(pg_num_type)
        else:
            sect_pr.append(pg_num_type)
    pg_num_type.set(qn('w:start'), str(body_start_page))

    # Augmentar el marge inferior del peu per no estar enganxat al límit del paper.
    # w:footer = distància des del límit inferior del paper fins al peu (en twips).
    # 567 twips ≈ 1 cm. El valor original de la plantilla és 0 (al límit).
    pg_mar = sect_pr.find(qn('w:pgMar'))
    if pg_mar is not None:
        pg_mar.set(qn('w:footer'), '567')

    footer  = body_section.footer
    para    = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_elem  = para._p

    # Netejar runs existents
    for r in p_elem.findall(qn('w:r')):
        p_elem.remove(r)

    # Centrar el paràgraf
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    existing_jc = pPr.find(qn('w:jc'))
    if existing_jc is not None:
        pPr.remove(existing_jc)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)

    def _field_run(child_elem):
        """Crea un <w:r> amb format Open Sans 10pt i l'element fill indicat."""
        r    = OxmlElement('w:r')
        rPr  = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
            rFonts.set(qn(attr), FONT_SANS)
        sz   = OxmlElement('w:sz');   sz.set(qn('w:val'),   '20')  # 10pt
        szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '20')
        rPr.extend([rFonts, sz, szCs])
        r.append(rPr)
        r.append(child_elem)
        return r

    fc_begin = OxmlElement('w:fldChar'); fc_begin.set(qn('w:fldCharType'), 'begin')
    instr    = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' PAGE '
    fc_end   = OxmlElement('w:fldChar'); fc_end.set(qn('w:fldCharType'), 'end')

    p_elem.append(_field_run(fc_begin))
    p_elem.append(_field_run(instr))
    p_elem.append(_field_run(fc_end))


# ─── Substitució de marcadors a la plantilla ─────────────────────────────────

def _replace_in_para(para, marker: str, value: str):
    """Substitueix el marcador en els runs del paràgraf preservant format."""
    for run in para.runs:
        if marker in run.text:
            run.text = run.text.replace(marker, value)


def _fix_header_alignment(doc, text_width_twips: int = 8504):
    """Reestructura les capçaleres que contenen {{EDICIO}} per tenir:
      - {{EDICIO}} alineat a l'esquerra
      - {{AUTOR}}  alineat a la dreta (via tab stop dret al marge dret)
    No afecta seccions sense {{EDICIO}} (portada, índex).
    """
    for section in doc.sections:
        hdr = section.header
        if hdr.is_linked_to_previous:
            continue
        for para in hdr.paragraphs:
            if '{{EDICIO}}' not in para.text:
                continue
            p_elem = para._p

            # 1. Afegir tab stop dret al pPr
            ppr = p_elem.find(qn('w:pPr'))
            if ppr is None:
                ppr = OxmlElement('w:pPr')
                p_elem.insert(0, ppr)
            # L'estil "Header" ja té un tab dret a 8504 twips.
            # Només cal netejar el tab centrat a 4252 perquè el tab caràcter
            # salti directament al tab dret en comptes del centrat.
            tabs = ppr.find(qn('w:tabs'))
            if tabs is None:
                tabs = OxmlElement('w:tabs')
                rpr_in_ppr = ppr.find(qn('w:rPr'))
                if rpr_in_ppr is not None:
                    rpr_in_ppr.addprevious(tabs)
                else:
                    ppr.append(tabs)
            for t in list(tabs):
                tabs.remove(t)
            clear_center = OxmlElement('w:tab')
            clear_center.set(qn('w:val'), 'clear')
            clear_center.set(qn('w:pos'), '4252')
            tabs.append(clear_center)

            # 2. Extreu rPr del primer run existent (per mantenir el format)
            existing_runs = p_elem.findall(qn('w:r'))
            rpr_src = None
            if existing_runs:
                rpr_src = existing_runs[0].find(qn('w:rPr'))
            for r in existing_runs:
                p_elem.remove(r)

            # 3. Reconstrueix: Run({{EDICIO}}) + tab + Run({{AUTOR}})
            def _make_run(text=None, tab=False):
                r = OxmlElement('w:r')
                if rpr_src is not None:
                    r.append(deepcopy(rpr_src))
                if tab:
                    r.append(OxmlElement('w:tab'))
                else:
                    t = OxmlElement('w:t')
                    t.text = text
                    r.append(t)
                return r

            p_elem.append(_make_run('{{EDICIO}}'))
            p_elem.append(_make_run(tab=True))
            p_elem.append(_make_run('{{AUTOR}}'))


def _fill_template_markers(doc, markers: dict):
    """Substitueix tots els marcadors al document (paràgrafs, taules, caps/peus)."""
    for para in doc.paragraphs:
        for mk, val in markers.items():
            _replace_in_para(para, mk, val)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for mk, val in markers.items():
                        _replace_in_para(para, mk, val)
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer):
            if hf:
                for para in hf.paragraphs:
                    for mk, val in markers.items():
                        _replace_in_para(para, mk, val)


def _fix_cover_labels(template_doc, doi: str):
    """Post-processa el format de la portada:
    1. Fa cursiva el label 'Title:' (anglès → cursiva com Keywords/Abstract).
    """
    for table in template_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        # "Title:" és anglès → cursiva (igual que Keywords/Abstract)
                        if run.text.startswith('Title'):
                            run.font.italic = True


def _handle_multiline_marker(doc, marker: str, paragraphs: list):
    """Substitueix un marcador per múltiples paràgrafs (sumari, abstract)."""
    if not paragraphs:
        _fill_template_markers(doc, {marker: ''})
        return
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if marker not in para.text:
                        continue
                    # Substituïm el text del marcador pel primer paràgraf
                    for run in para.runs:
                        if marker in run.text:
                            run.text = run.text.replace(marker, paragraphs[0])
                            break
                    # Inserim els paràgrafs addicionals just després
                    ref_p = para._p
                    for extra in reversed(paragraphs[1:]):
                        new_p = OxmlElement('w:p')
                        # Copiem propietats del paràgraf original (alineació, etc.)
                        orig_ppr = ref_p.find(qn('w:pPr'))
                        if orig_ppr is not None:
                            new_p.append(deepcopy(orig_ppr))
                        new_r = OxmlElement('w:r')
                        # Copiem el format del run original
                        orig_rpr = ref_p.find('.//' + qn('w:rPr'))
                        if orig_rpr is not None:
                            new_r.append(deepcopy(orig_rpr))
                        new_t = OxmlElement('w:t')
                        new_t.text = extra
                        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                        new_r.append(new_t)
                        new_p.append(new_r)
                        ref_p.addnext(new_p)
                    return


def _copy_numbering_from_article(orig_doc, tmpl_doc):
    """Substitueix numbering.xml de la plantilla pel de l'article original.
    Canvia tots els nivells amb format upperLetter + sufix ) a lowerLetter + .
    Necessari perquè w:numPr de l'índex copiat apuntin a numIds vàlids.
    """
    try:
        orig_nb = orig_doc.part.numbering_part
        tmpl_nb = tmpl_doc.part.numbering_part
        tmpl_elem = tmpl_nb._element
        for child in list(tmpl_elem):
            tmpl_elem.remove(child)
        for child in orig_nb._element:
            tmpl_elem.append(deepcopy(child))
        # A) → a.  en tots els lvl amb upperLetter i sufix )
        # I elimina el desplaçament base de cada abstractNum: resta el w:left
        # de ilvl=0 a tots els nivells perquè N1 quedi alineat amb el marge.
        for an in tmpl_elem.findall(qn('w:abstractNum')):
            # Troba w:left i w:hanging del ilvl=0
            base_left = base_hanging = 0
            for lvl in an.findall(qn('w:lvl')):
                if lvl.get(qn('w:ilvl')) == '0':
                    ppr = lvl.find(qn('w:pPr'))
                    if ppr is not None:
                        ind = ppr.find(qn('w:ind'))
                        if ind is not None:
                            base_left    = int(ind.get(qn('w:left'),    '0'))
                            base_hanging = int(ind.get(qn('w:hanging'), '0'))
                    break
            # offset = posició del número en ilvl=0 (el que volem moure a 0)
            offset = base_left - base_hanging
            if offset <= 0:
                continue
            # Resta offset a w:left de cada nivell (w:hanging queda intacte)
            for lvl in an.findall(qn('w:lvl')):
                ppr = lvl.find(qn('w:pPr'))
                if ppr is None:
                    continue
                ind = ppr.find(qn('w:ind'))
                if ind is None:
                    continue
                left    = int(ind.get(qn('w:left'),    '0'))
                hanging = int(ind.get(qn('w:hanging'), '0'))
                new_left = max(hanging, left - offset)
                ind.set(qn('w:left'), str(new_left))

        for lvl in tmpl_elem.iter(qn('w:lvl')):
            nf = lvl.find(qn('w:numFmt'))
            lt = lvl.find(qn('w:lvlText'))
            if nf is None or lt is None:
                continue
            if (nf.get(qn('w:val')) == 'upperLetter'
                    and lt.get(qn('w:val'), '').endswith(')')):
                nf.set(qn('w:val'), 'lowerLetter')
                lt.set(qn('w:val'), lt.get(qn('w:val'))[:-1] + '.')
    except Exception as e:
        print(f"  [!] No s'ha pogut copiar numbering: {e}")


def _build_numformat_map(tmpl_doc) -> dict:
    """Retorna {(numId_str, ilvl_str): numFmt_str} de la plantilla."""
    result = {}
    try:
        nb_elem = tmpl_doc.part.numbering_part._element
        abstract_map = {}
        for an in nb_elem.findall(qn('w:abstractNum')):
            an_id = an.get(qn('w:abstractNumId'))
            lvl_fmts = {}
            for lvl in an.findall(qn('w:lvl')):
                il = lvl.get(qn('w:ilvl'))
                nf = lvl.find(qn('w:numFmt'))
                if nf is not None:
                    lvl_fmts[il] = nf.get(qn('w:val'), '')
            abstract_map[an_id] = lvl_fmts
        for num in nb_elem.findall(qn('w:num')):
            num_id = num.get(qn('w:numId'))
            an_ref = num.find(qn('w:abstractNumId'))
            if an_ref is None:
                continue
            lvl_fmts = abstract_map.get(an_ref.get(qn('w:val')), {})
            for il, fmt in lvl_fmts.items():
                result[(num_id, il)] = fmt
    except Exception:
        pass
    return result


def _copy_index_to_template(doc, index_paras: list, levels: list = None,
                            numformat_map: dict = None):
    """Substitueix {{INDEX}} per còpies exactes dels paràgrafs d'índex originals.
    Preserva estructura original: números, tabulats i indentats via w:numPr.
    Aplica PT Serif 9pt.
    Cursiva NOMÉS per ilvl=0 amb numFmt decimal (1., 2., 3.).
    Entrades de lletra (a., b., c.) i subnivells: sense cursiva.
    """
    def _apply_style(p_elem, explicit_level=None):
        ppr      = p_elem.find(qn('w:pPr'))
        ilvl_str = '0'
        num_id   = None
        if ppr is not None:
            numpr = ppr.find(qn('w:numPr'))
            if numpr is not None:
                il = numpr.find(qn('w:ilvl'))
                ni = numpr.find(qn('w:numId'))
                if il is not None:
                    ilvl_str = il.get(qn('w:val'), '0')
                if ni is not None:
                    num_id = ni.get(qn('w:val'))

        num_fmt = (numformat_map or {}).get((num_id, ilvl_str), '')
        if explicit_level is not None and num_id is None:
            # Paràgraf construït de zero (sense w:numPr): usar el nivell explícit
            is_n1 = (explicit_level == 1)
        else:
            # Cursiva: ilvl=0 + decimal (1., 2., 3.) → sí. Resta → no.
            is_n1 = (ilvl_str == '0' and num_fmt == 'decimal')

        # Alineació esquerra + interlineat simple al pPr
        if ppr is None:
            ppr = OxmlElement('w:pPr')
            p_elem.insert(0, ppr)
        jc = ppr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            ppr.append(jc)
        jc.set(qn('w:val'), 'left')
        sp = ppr.find(qn('w:spacing'))
        if sp is None:
            sp = OxmlElement('w:spacing')
            ppr.append(sp)
        sp.set(qn('w:line'), '240')
        sp.set(qn('w:lineRule'), 'auto')
        sp.set(qn('w:after'), '0')
        sp.set(qn('w:before'), '0')

        # Entrades de lletra (a., b., c.): forçar indentació un nivell per sota de 2.x.x
        # Les seves numIds tenen ilvl=0 però han d'aparèixer com a nivell 3.
        # Nivells decimal: step=720 twips, hanging=360.
        # ilvl=0 → left=360 | ilvl=1 → left=1080 | ilvl=2 → left=1800 | lletra → left=2520
        if num_fmt == 'lowerLetter':
            ind = ppr.find(qn('w:ind'))
            if ind is None:
                ind = OxmlElement('w:ind')
                ppr.append(ind)
            ind.set(qn('w:left'),    '1820')
            ind.set(qn('w:hanging'), '360')

        for rpr in p_elem.iter(qn('w:rPr')):
            rf = rpr.find(qn('w:rFonts'))
            if rf is None:
                rf = OxmlElement('w:rFonts')
                rpr.insert(0, rf)
            rf.set(qn('w:ascii'), FONT_SERIF)
            rf.set(qn('w:hAnsi'), FONT_SERIF)
            rf.set(qn('w:cs'),    FONT_SERIF)
            for theme_attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme', 'w:eastAsiaTheme'):
                if rf.get(qn(theme_attr)) is not None:
                    del rf.attrib[qn(theme_attr)]
            for tag in ('w:sz', 'w:szCs'):
                el = rpr.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    rpr.append(el)
                el.set(qn('w:val'), '18')  # 9pt
            i_el = rpr.find(qn('w:i'))
            if is_n1:
                if i_el is None:
                    rpr.append(OxmlElement('w:i'))
            else:
                if i_el is not None:
                    rpr.remove(i_el)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if '{{INDEX}}' not in para.text:
                        continue
                    ref_p  = para._p
                    parent = ref_p.getparent()
                    lvl_list = levels or []
                    for orig_p, lvl in zip(reversed(index_paras),
                                           reversed(lvl_list) if lvl_list
                                           else [None] * len(index_paras)):
                        new_p = deepcopy(orig_p)
                        _apply_style(new_p, explicit_level=lvl)
                        ref_p.addnext(new_p)
                    parent.remove(ref_p)
                    return  # només hi ha un {{INDEX}}


def _append_body_to_template(template_doc, article_doc, body_start: int, body_end: int):
    """Substitueix el cos demo de la plantilla pel cos de l'article (ja corregit)."""
    body = template_doc.element.body

    # Trobar el segon paràgraf amb sectPr (salt de secció després de l'índex)
    sect_pr_paras = [
        child for child in body
        if child.tag == qn('w:p') and child.find('.//' + qn('w:sectPr')) is not None
    ]
    second_sect_para = sect_pr_paras[1] if len(sect_pr_paras) >= 2 else None

    # Esborrar tots els elements entre el segon sectPr i el sectPr final
    if second_sect_para is not None:
        removing = False
        to_remove = []
        for child in list(body):
            if removing:
                if child.tag == qn('w:sectPr'):
                    break
                to_remove.append(child)
            elif child is second_sect_para:
                removing = True
        for elem in to_remove:
            body.remove(elem)

    # Inserir els paràgrafs de l'article abans del sectPr final
    final_sect_pr = body.find(qn('w:sectPr'))
    paras = article_doc.paragraphs
    for i in range(body_start, min(body_end, len(paras))):
        elem = deepcopy(paras[i]._element)
        if final_sect_pr is not None:
            final_sect_pr.addprevious(elem)
        else:
            body.append(elem)


# ─── Informe ─────────────────────────────────────────────────────────────────
class Report:
    def __init__(self, filename: str):
        self.filename = filename
        self.applied  = []
        self.alerts   = []
        self.pending  = [
            "Número d'edició InDret (assignat per la redacció)",
            "Dates de recepció i acceptació (afegides per InDret)",
            "Paginació final (a càrrec de l'equip d'edició)",
            "Encapçalament: nom autor (dreta) / nº InDret (esquerra) — Open Sans 8 negrita",
        ]

    def ok(self,   msg: str): self.applied.append(msg)
    def warn(self, msg: str): self.alerts.append(msg)

    def to_markdown(self) -> str:
        now = datetime.now().strftime("%Y-%m-%d %H:%M")
        lines = [f"# Informe de correcció — {self.filename} — {now}", ""]
        lines += ["## Canvis aplicats automàticament", ""]
        for i in self.applied: lines.append(f"- [✓] {i}")
        lines += ["", "## Alertes que requereixen revisió manual", ""]
        if self.alerts:
            for i in self.alerts: lines.append(f"- [!] {i}")
        else:
            lines.append("- Cap alerta detectada.")
        lines += ["", "## Elements incomplets (afegir per InDret)", ""]
        for i in self.pending: lines.append(f"- [ ] {i}")
        return "\n".join(lines)


# ─── Corrector principal ──────────────────────────────────────────────────────
class InDretCorrector:

    def __init__(self, input_path: str):
        self.path    = Path(input_path)
        self.doc     = Document(input_path)
        self.report  = Report(self.path.name)
        self._headings: dict[str, str] = {}  # text → bookmark_name

    def run(self) -> tuple[str, str]:
        print("  [1/5] Correccions de text (espais, cometes)")
        self._phase1_text()
        print("  [2/5] Estils tipogràfics i espaiat")
        self._phase2_styles()
        print("  [3/5] Versaletes (cos, notes al peu, bibliografía)")
        self._phase3_small_caps()
        print("  [4/5] Índex (sagnats, hiperenllaços)")
        self._phase4_index()
        print("  [5/5] Verificacions (jurisprudència, op. cit.)")
        self._phase5_checks()
        return self._save()

    # ── Fase 1: correccions de text ──────────────────────────────────────────
    def _phase1_text(self):
        total_sp = total_qt = 0
        for para in self.doc.paragraphs:
            for run in para.runs:
                new_t, n_sp, n_qt = self._fix_text(run.text)
                if new_t != run.text:
                    run.text = new_t
                total_sp += n_sp
                total_qt += n_qt
        if total_sp: self.report.ok(f"Espais dobles eliminats: {total_sp} instàncies")
        if total_qt: self.report.ok(f"Cometes angleses → «»: {total_qt} substitucions")
        if not total_sp and not total_qt:
            self.report.ok("Text: cap espai doble ni cometa incorrecta")

    @staticmethod
    def _fix_text(text: str) -> tuple[str, int, int]:
        new, n_sp = RE_DOUBLE_SPACE.subn(' ', text)
        n_qt = 0
        if '"' in new:
            before = new.count('"')
            new    = RE_QUOT_OPEN.sub(r'«\1', new)
            new    = RE_QUOT_CLOSE.sub(r'\1»', new)
            n_qt   = max(0, before - new.count('"'))
        return new, n_sp, n_qt

    # ── Fase 2: estils tipogràfics ────────────────────────────────────────────
    def _phase2_styles(self):
        in_bib        = False
        in_index      = False
        cnt_body      = cnt_h = 0
        h1_num_levels = set()  # (numId, ilvl) dels h1 amb w:numPr

        for para in self.doc.paragraphs:
            text  = para.text.strip()
            tl    = text.lower()
            ptype = classify_para(para)

            if ptype == 'bib_keyword':
                in_bib   = True
                in_index = False
                # Aplica format h1 al títol de bibliografía (Open Sans 11 bold)
                for r in para.runs:
                    set_run_font(r, FONT_SANS, Pt(11), bold=True, italic=False)
                set_line_spacing(para, 1.0)
                para.paragraph_format.space_before = Pt(8)
                para.paragraph_format.space_after  = Pt(18)
                continue
            if ptype == 'index_keyword':
                in_index = True
                continue
            # Fi de la zona d'índex: primera línia en blanc després del títol
            if in_index and not text:
                in_index = False

            if ptype == 'empty':
                continue

            # ── Títols de l'índex: s'apliquen a _phase4_index, aquí saltem
            if in_index:
                continue

            if ptype == 'h1':
                for r in para.runs:
                    set_run_font(r, FONT_SANS, Pt(11), bold=True, italic=False)
                # Recollir numId/ilvl per corregir la font al numbering.xml
                pPr_el = para._p.find(qn('w:pPr'))
                if pPr_el is not None:
                    numPr_el = pPr_el.find(qn('w:numPr'))
                    if numPr_el is not None:
                        ilvl_el  = numPr_el.find(qn('w:ilvl'))
                        numId_el = numPr_el.find(qn('w:numId'))
                        if ilvl_el is not None and numId_el is not None:
                            h1_num_levels.add((
                                numId_el.get(qn('w:val')),
                                ilvl_el.get(qn('w:val')),
                            ))
                set_line_spacing(para, 1.0)
                para.paragraph_format.space_before = Pt(8)
                para.paragraph_format.space_after  = Pt(18)
                cnt_h += 1

            elif ptype == 'h2':
                for r in para.runs:
                    set_run_font(r, FONT_SERIF, Pt(10), bold=True, italic=False)
                set_line_spacing(para, 1.0)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after  = Pt(3)
                cnt_h += 1

            elif ptype in ('h3', 'h4'):
                for r in para.runs:
                    set_run_font(r, FONT_SERIF, Pt(10), bold=False, italic=True)
                set_line_spacing(para, 1.0)
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after  = Pt(2)
                cnt_h += 1

            else:  # body / bibliography
                pf = para.paragraph_format
                pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pf.first_line_indent = Pt(0)   # sense sagnat de primera línia
                for r in para.runs:
                    b  = bool(r.bold)
                    it = bool(r.italic)
                    set_run_font(r, FONT_SERIF, Pt(10), bold=b, italic=it)
                set_line_spacing(para, 1.1)
                if not pf.space_after or pf.space_after.pt == 0:
                    pf.space_after = Pt(6)
                cnt_body += 1

        # Font del número auto-generat (w:numPr) als títols N1
        if h1_num_levels:
            self._fix_numpr_font(h1_num_levels)

        # Notes al peu: font PT Serif 8,5
        self._apply_footnote_font()

        self.report.ok(f"Fonts aplicades: {cnt_h} títols, {cnt_body} paràgrafs de cos")
        self.report.ok("Cos del text → PT Serif 10, interlineat 1,1, space_after 6pt")
        self.report.ok("Títols N1 → Open Sans 11 negrita | N2 → PT Serif 10 negrita | N3/N4 → PT Serif 10 cursiva")

    def _fix_numpr_font(self, num_levels: set):
        """Actualitza w:abstractNum/w:lvl/w:rPr al numbering.xml per als nivells
        indicats (numId, ilvl), forçant Open Sans 11pt bold (font del label N1)."""
        try:
            num_part = getattr(self.doc.part, 'numbering_part', None)
            if num_part is None:
                return
        except Exception:
            return
        num_elem = num_part._element

        # numId → abstractNumId
        num_to_abstract = {}
        for num_el in num_elem.findall(qn('w:num')):
            nid = num_el.get(qn('w:numId'))
            abs_el = num_el.find(qn('w:abstractNumId'))
            if abs_el is not None:
                num_to_abstract[nid] = abs_el.get(qn('w:val'))

        for num_id_str, ilvl_str in num_levels:
            abs_num_id = num_to_abstract.get(num_id_str)
            if abs_num_id is None:
                continue
            for abs_num in num_elem.findall(qn('w:abstractNum')):
                if abs_num.get(qn('w:abstractNumId')) != abs_num_id:
                    continue
                for lvl in abs_num.findall(qn('w:lvl')):
                    if lvl.get(qn('w:ilvl')) != ilvl_str:
                        continue
                    rPr = lvl.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        lvl.append(rPr)
                    rf = rPr.find(qn('w:rFonts'))
                    if rf is None:
                        rf = OxmlElement('w:rFonts')
                        rPr.insert(0, rf)
                    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
                        rf.set(qn(attr), FONT_SANS)
                    for theme_attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme', 'w:eastAsiaTheme'):
                        if rf.get(qn(theme_attr)) is not None:
                            del rf.attrib[qn(theme_attr)]
                    for tag in ('w:sz', 'w:szCs'):
                        el = rPr.find(qn(tag))
                        if el is None:
                            el = OxmlElement(tag)
                            rPr.append(el)
                        el.set(qn('w:val'), '22')  # 11pt = 22 half-points
                    if rPr.find(qn('w:b')) is None:
                        rPr.append(OxmlElement('w:b'))

    def _apply_footnote_font(self):
        try:
            fn_part = getattr(self.doc.part, 'footnotes_part', None)
            if not fn_part: return
            n = 0
            for fn in fn_part._element.findall('.//' + qn('w:footnote')):
                fn_id = int(fn.get(qn('w:id'), 0))
                if fn_id <= 0: continue
                for r in fn.findall('.//' + qn('w:r')):
                    rpr = r.find(qn('w:rPr'))
                    if rpr is None:
                        rpr = OxmlElement('w:rPr')
                        r.insert(0, rpr)
                    rf = rpr.find(qn('w:rFonts'))
                    if rf is None:
                        rf = OxmlElement('w:rFonts')
                        rpr.insert(0, rf)
                    rf.set(qn('w:ascii'), FONT_SERIF)
                    rf.set(qn('w:hAnsi'), FONT_SERIF)
                    for tag in ('w:sz', 'w:szCs'):
                        el = rpr.find(qn(tag))
                        if el is None:
                            el = OxmlElement(tag)
                            rpr.append(el)
                        el.set(qn('w:val'), '17')   # 8.5pt = 17 half-points
                n += 1
            if n:
                self.report.ok(f"Notes al peu → PT Serif 8,5 ({n} notes)")
        except Exception as e:
            self.report.warn(f"No s'han pogut processar les notes al peu: {e}")

    # ── Fase 3: versaletes ────────────────────────────────────────────────────
    def _phase3_small_caps(self):
        # 3a. Cos del text i bibliografía
        caps_body = self._small_caps_paragraphs()
        # 3b. Notes al peu
        caps_fn   = self._small_caps_footnotes()
        if caps_body + caps_fn:
            self.report.ok(
                f"Versaletes aplicades: {caps_body} cognoms al cos/bibliografía, "
                f"{caps_fn} a notes al peu"
            )
        else:
            self.report.ok("Versaletes: no s'han detectat cognoms en MAJÚSCULES")

    def _small_caps_paragraphs(self) -> int:
        count = 0
        for para in self.doc.paragraphs:
            if not para.text.strip():
                continue
            for run in para.runs:
                if is_likely_surname(run.text):
                    run.text = run.text.strip().title()
                    set_small_caps_xml(run._r)
                    count += 1
        return count

    def _small_caps_footnotes(self) -> int:
        count = 0
        try:
            fn_part = getattr(self.doc.part, 'footnotes_part', None)
            if not fn_part: return 0
            for fn in fn_part._element.findall('.//' + qn('w:footnote')):
                fn_id = int(fn.get(qn('w:id'), 0))
                if fn_id <= 0: continue
                for r in fn.findall('.//' + qn('w:r')):
                    t = r.find(qn('w:t'))
                    if t is not None and t.text and is_likely_surname(t.text):
                        t.text = t.text.strip().title()
                        set_small_caps_xml(r)
                        count += 1
        except Exception:
            pass
        return count

    # ── Fase 4: índex ─────────────────────────────────────────────────────────
    def _phase4_index(self):
        paras = self.doc.paragraphs

        # Pas 1: recollir tots els títols del cos → mapa text→bookmark
        bm_id = self._get_max_existing_bookmark_id() + 1
        for para in paras:
            text  = para.text.strip()
            level = get_heading_level(text)
            if level > 0 and text not in self._headings:
                bm_name = f"bm_indret_{bm_id}"
                self._headings[text] = bm_name
                add_bookmark(para, bm_id, bm_name)
                bm_id += 1

        # Pas 2: trobar secció d'índex
        index_start = next(
            (i for i, p in enumerate(paras) if p.text.strip().lower() in INDEX_KEYWORDS),
            -1
        )
        if index_start < 0:
            self.report.warn("No s'ha detectat la secció d'índex (títol 'Índice'/'Sumario')")
            return

        # Fi de l'índex: primer paràgraf llarg (cos del text)
        index_end = next(
            (i for i in range(index_start + 1, len(paras))
             if paras[i].text.strip() and len(paras[i].text.strip()) > 120),
            len(paras)
        )

        indent_cnt = link_cnt = 0

        for i in range(index_start + 1, index_end):
            para  = paras[i]
            text  = para.text.strip()
            if not text: continue
            level = get_heading_level(text)
            if level == 0: continue

            # Sagnat per nivell
            para.paragraph_format.left_indent = INDEX_INDENT.get(level, Cm(0.0))
            indent_cnt += 1

            # Tipografia d'índex: N1 → PT Serif 9 cursiva negrita; N2+ → PT Serif 9 normal
            for run in para.runs:
                if level == 1:
                    set_run_font(run, FONT_SERIF, Pt(9), bold=True, italic=True)
                else:
                    set_run_font(run, FONT_SERIF, Pt(9), bold=False, italic=False)

            # Hiperenllaç intern si el text coincideix amb un títol del cos
            if text in self._headings:
                add_internal_hyperlink(para, self._headings[text])
                link_cnt += 1

        if indent_cnt:
            self.report.ok(f"Índex: sagnats per nivell aplicats ({indent_cnt} entrades)")
        if link_cnt:
            self.report.ok(f"Índex: {link_cnt} hiperenllaços interns creats")
        elif indent_cnt:
            self.report.warn(
                "Índex: no s'han pogut crear hiperenllaços — el text de les entrades "
                "no coincideix exactament amb els títols del cos del text"
            )

    def _get_max_existing_bookmark_id(self) -> int:
        """Retorna el màxim ID de bookmark ja existent al document."""
        max_id = 0
        for elem in self.doc.element.iter(qn('w:bookmarkStart')):
            try:
                max_id = max(max_id, int(elem.get(qn('w:id'), 0)))
            except (ValueError, TypeError):
                pass
        return max_id

    # ── Fase 5: verificacions ─────────────────────────────────────────────────
    def _phase5_checks(self):
        in_bib = False
        has_abstract = has_summary = has_keywords = False

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            tl   = text.lower()

            if tl in BIB_KEYWORDS:
                in_bib = True; continue
            if not text: continue

            if 'abstract'   in tl: has_abstract  = True
            if any(k in tl for k in ('sumario', 'resumen', 'sumari', 'resum')):
                has_summary = True
            if any(k in tl for k in ('palabras clave', 'keywords', 'paraules clau')):
                has_keywords = True

            # Jurisprudència al cos (fora de bibliografía)
            if not in_bib and RE_JURIS.search(text):
                if not RE_ECLI.search(text) and not RE_ROJ.search(text):
                    m = RE_JURIS.search(text)
                    self.report.warn(
                        f"Paràgraf ~{i+1}: cita de jurisprudència ({m.group()}) "
                        f"sense referència ECLI ni Roj"
                    )

        # Fallback: escanejar les taules (portada ja formatada)
        if not has_abstract or not has_summary or not has_keywords:
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            tl = p.text.strip().lower()
                            if not tl:
                                continue
                            if 'abstract' in tl:
                                has_abstract = True
                            if any(k in tl for k in ('sumario', 'resumen', 'sumari', 'resum')):
                                has_summary = True
                            if any(k in tl for k in ('palabras clave', 'keywords', 'paraules clau')):
                                has_keywords = True

        # op. cit. a les notes al peu
        self._check_opcit_footnotes()

        if not has_abstract:  self.report.warn("No s'ha detectat abstract en anglès (requerit per portada)")
        if not has_summary:   self.report.warn("No s'ha detectat sumari/resum")
        if not has_keywords:  self.report.warn("No s'han detectat paraules clau / keywords")

    def _check_opcit_footnotes(self):
        try:
            fn_part = getattr(self.doc.part, 'footnotes_part', None)
            if not fn_part: return
            for fn in fn_part._element.findall('.//' + qn('w:footnote')):
                fn_id = int(fn.get(qn('w:id'), 0))
                if fn_id <= 0: continue
                fn_text = ''.join(
                    (t.text or '') for t in fn.findall('.//' + qn('w:t'))
                )
                if RE_OP_CIT.search(fn_text):
                    self.report.warn(
                        f"Nota {fn_id}: conté «op. cit.» — verificar si és la primera "
                        f"citació d'aquesta obra (§9: la primera cita ha de ser completa)"
                    )
        except Exception:
            pass

    def _get_article_footnotes_xml(self) -> tuple[bytes | None, bytes | None]:
        """Retorna (footnotes_xml processat, footnotes_rels) o (None, None)."""
        raw = None
        fn_rels = None
        fn_part = getattr(self.doc.part, 'footnotes_part', None)
        if fn_part is not None:
            raw = _lxml_et.tostring(fn_part._element, xml_declaration=True,
                                    encoding='UTF-8', standalone=True)
        try:
            with zipfile.ZipFile(str(self.path), 'r') as z:
                if raw is None and 'word/footnotes.xml' in z.namelist():
                    raw = z.read('word/footnotes.xml')
                if 'word/_rels/footnotes.xml.rels' in z.namelist():
                    fn_rels = z.read('word/_rels/footnotes.xml.rels')
        except Exception:
            pass
        if raw is None:
            return None, None
        return self._process_footnotes_xml(raw), fn_rels

    def _process_footnotes_xml(self, xml_bytes: bytes) -> bytes:
        """Aplica correccions al XML de footnotes:
        - Elimina cursiva del número de nota (w:footnoteRef)
        - Afegeix espaiat entre notes (space_after al darrer paràgraf de cada nota)
        """
        W_FN      = qn('w:footnote')
        W_ID      = qn('w:id')
        W_P       = qn('w:p')
        W_R       = qn('w:r')
        W_RPR     = qn('w:rPr')
        W_I       = qn('w:i')
        W_FNREF   = qn('w:footnoteRef')
        W_PPR     = qn('w:pPr')
        W_SPACING = qn('w:spacing')

        root = _lxml_et.fromstring(xml_bytes)
        for fn in root:
            if fn.tag != W_FN:
                continue
            try:
                fn_id = int(fn.get(W_ID, 0))
            except (ValueError, TypeError):
                continue
            if fn_id < 1:
                continue

            paras = fn.findall(W_P)

            # Normalitza estil i espaiat de cada paràgraf de la nota:
            # - Substitueix qualsevol estil que no existeixi al template per FootnoteText
            # - Força spacing: line=240/auto, before=0, after=0
            KEEP_STYLES = {'FootnoteText', 'Endnotetext'}
            for p in paras:
                pPr = p.find(W_PPR)
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    p.insert(0, pPr)
                ps_el = pPr.find(qn('w:pStyle'))
                if ps_el is None:
                    ps_el = OxmlElement('w:pStyle')
                    pPr.insert(0, ps_el)
                if ps_el.get(qn('w:val'), '') not in KEEP_STYLES:
                    ps_el.set(qn('w:val'), 'FootnoteText')
                sp = pPr.find(W_SPACING)
                if sp is None:
                    sp = OxmlElement('w:spacing')
                    pPr.append(sp)
                sp.set(qn('w:line'), '212')       # ~1.0× per 8.5pt
                sp.set(qn('w:lineRule'), 'auto')
                sp.set(qn('w:before'), '0')
                sp.set(qn('w:after'), '0')
            # Petit espaiat després de l'última línia de cada nota (separa de la següent)
            if paras:
                last_pPr = paras[-1].find(W_PPR)
                last_sp = last_pPr.find(W_SPACING)
                last_sp.set(qn('w:after'), '100')  # 5pt entre notes

            for r in fn.iter(W_R):
                is_ref = r.find(W_FNREF) is not None

                # Font PT Serif 8,5pt a tots els runs (inclòs el del número)
                rPr = r.find(W_RPR)
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r.insert(0, rPr)
                rf = rPr.find(qn('w:rFonts'))
                if rf is None:
                    rf = OxmlElement('w:rFonts')
                    rPr.insert(0, rf)
                for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
                    rf.set(qn(attr), FONT_SERIF)
                for theme_attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme', 'w:eastAsiaTheme'):
                    if rf.get(qn(theme_attr)) is not None:
                        del rf.attrib[qn(theme_attr)]
                for sz_tag in ('w:sz', 'w:szCs'):
                    el = rPr.find(qn(sz_tag))
                    if el is None:
                        el = OxmlElement(sz_tag)
                        rPr.append(el)
                    el.set(qn('w:val'), '17')  # 8,5pt = 17 half-points

                if is_ref:
                    # Número de nota: substituïm el rPr per un de net
                    # (superíndex, sense cursiva heretada de cap estil)
                    r.remove(rPr)
                    new_rPr = OxmlElement('w:rPr')
                    for wtag in ('w:i', 'w:iCs'):
                        el = OxmlElement(wtag)
                        el.set(qn('w:val'), '0')
                        new_rPr.append(el)
                    vert = OxmlElement('w:vertAlign')
                    vert.set(qn('w:val'), 'superscript')
                    new_rPr.append(vert)
                    r.insert(0, new_rPr)


        return _lxml_et.tostring(root, xml_declaration=True,
                                 encoding='UTF-8', standalone=True)

    def _number_h1_headings(self, body_start: int, body_end: int) -> dict:
        """Afegeix numeració (1., 2., ...) als títols N1 que no en tenen.

        Retorna {text_original: text_nou} per actualitzar les entrades d'índex.
        Si algun títol N1 ja té numeració, s'assumeix que l'article és consistent i
        no es fa cap canvi.
        """
        paras = self.doc.paragraphs
        n     = min(body_end, len(paras))

        h1_paras = [
            (i, paras[i])
            for i in range(body_start, n)
            if classify_para(paras[i]) == 'h1'
            and paras[i].text.strip().lower() not in NO_NUMBER_HEADINGS
        ]
        if not h1_paras:
            return {}

        # Si algun títol ja té numeració → l'article és consistent, no tocar
        if any(re.match(r'^\d+\.', p.text.strip()) for _, p in h1_paras):
            return {}

        remapping = {}
        for counter, (_, para) in enumerate(h1_paras, start=1):
            old_text = para.text.strip()
            prefix   = f"{counter}. "
            for run in para.runs:
                if run.text:
                    run.text = prefix + run.text
                    break
            remapping[old_text] = prefix + old_text

        if remapping:
            self.report.ok(f"Títols N1 numerats automàticament: {len(remapping)} títols")
        return remapping

    # ── Flux basat en plantilla ───────────────────────────────────────────────
    def template_run(self, plantilla_path: str, edicio: str = '', doi: str = '',
                     autor: str = '', recepcio: str = '', acceptacio: str = '',
                     pagina_inici: int = 1) -> tuple[str, str]:
        """Flux alternatiu: corregeix l'article i l'incorpora a la plantilla."""
        print("  [1/5] Correccions de text (espais, cometes)")
        self._phase1_text()
        print("  [2/5] Estils tipogràfics i espaiat")
        self._phase2_styles()
        print("  [3/5] Versaletes (cos, notes al peu, bibliografía)")
        self._phase3_small_caps()

        print("  [4/5] Extracció de metadades de l'article")
        extractor = MetadataExtractor()
        data = extractor.extract(self.doc)

        # Avisos per camps no trobats
        if not data['titol']:
            self.report.warn("No s'ha pogut extreure el títol de l'article")
        if autor:
            data['autor'] = autor  # paràmetre manual té prioritat
        if not data['autor']:
            data['autor'] = 'Nombre Autor'
            self.report.warn("Autor no detectat — s'ha inserit 'Nombre Autor' com a placeholder")
        if not data['sumari']:
            self.report.warn("Sumari/Resum no detectat — verificar secció RESUMEN/SUMARIO")
        if not data['abstract']:
            self.report.warn("Abstract no detectat — verificar secció ABSTRACT")
        if not data['index_entries']:
            self.report.warn("Índex no detectat — verificar secció ÍNDICE/ÍNDEX")
        if not edicio:
            self.report.warn("Número d'edició no proporcionat — {{EDICIO}} quedarà buit")
        if recepcio:
            self.report.ok(f"Data de recepció: {recepcio}")
            self.report.pending = [p for p in self.report.pending
                                   if 'recepci' not in p.lower()]
        if acceptacio:
            self.report.ok(f"Data d'acceptació: {acceptacio}")
            self.report.pending = [p for p in self.report.pending
                                   if 'acceptaci' not in p.lower() and 'recepci' not in p.lower()]

        self.report.ok(f"Títol extret: {data['titol']}")
        if data['titol_en']:
            self.report.ok(f"Títol anglès: {data['titol_en']}")
        if data['index_entries']:
            self.report.ok(f"Índex: {len(data['index_entries'])} entrades detectades")
        body_n = data['body_end_idx'] - data['body_start_idx']
        self.report.ok(f"Cos de l'article: {body_n} paràgrafs (des de [{data['body_start_idx']}])")


        # Notes al peu: obtenir XML processat (s'injectarà al ZIP de sortida)
        article_fn_xml, article_fn_rels = self._get_article_footnotes_xml()
        if article_fn_xml is not None:
            self.report.ok("Notes al peu detectades — s'inclouen al document final")

        print("  [5/5] Verificacions + ompliment de plantilla")
        self._phase5_checks()

        # Carregar plantilla i substituir marcadors simples (una línia)
        template_doc = Document(plantilla_path)
        # Reestructura capçaleres del cos amb tab stop dret abans d'omplir marcadors
        _fix_header_alignment(template_doc)
        autor_display = data['autor'] if data['autor'] else 'Nombre Autor'
        edicio_display = f"InDret {edicio.replace('/', '.')}".strip() if edicio else "InDret"
        markers = {
            '{{TITOL}}':        data['titol'],
            '{{SUBTITOL}}':     data['subtitol'],
            '{{AUTOR}}':        autor_display,
            '{{ORGANITZACIO}}': data['organitzacio'],
            '{{EDICIO}}':       edicio_display,
            '{{TITOL_EN}}':     data['titol_en'],
            '{{PARAULES_CLAU}}': data['paraules_clau'],
            '{{KEYWORDS}}':     data['keywords'],
            '{{DOI}}':          doi,
            '{{DATA-RECEP}}':   recepcio,
            '{{DATA-ACCEPT}}':  acceptacio,
        }
        _fill_template_markers(template_doc, markers)
        _fix_cover_labels(template_doc, doi)

        # Marcadors multiparàgraf (sumari i abstract)
        _handle_multiline_marker(template_doc, '{{SUMARI}}',   data['sumari'])
        _handle_multiline_marker(template_doc, '{{ABSTRACT}}', data['abstract'])

        # Omplir l'índex: còpia directa dels paràgrafs originals (preserva format)
        if data['index_paras']:
            # Copia numbering.xml de l'original per preservar w:numPr de l'índex
            _copy_numbering_from_article(self.doc, template_doc)
            numfmt_map = _build_numformat_map(template_doc)
            levels = [lvl for _, lvl in data['index_entries']]
            _copy_index_to_template(template_doc, data['index_paras'], levels, numfmt_map)
        else:
            _fill_template_markers(template_doc, {'{{INDEX}}': ''})

        # Afegir el cos de l'article (ja corregit) a continuació de la plantilla
        _append_body_to_template(
            template_doc, self.doc,
            data['body_start_idx'], data['body_end_idx']
        )

        # Numeració de pàgines: peu centrat Open Sans 10pt, des de la secció del cos
        _add_page_numbers(template_doc, body_start_page=pagina_inici)
        self.report.ok(f"Numeració de pàgines afegida (peu centrat, Open Sans 10pt, inici pàg. {pagina_inici})")

        out_doc, out_report = self._save_doc(template_doc)

        # Injectar notes al peu al ZIP de sortida
        if article_fn_xml is not None:
            try:
                _inject_footnotes(out_doc, article_fn_xml, article_fn_rels)
            except Exception as e:
                self.report.warn(f"No s'han pogut injectar les notes al peu: {e}")

        return out_doc, out_report

    # ── Desar ─────────────────────────────────────────────────────────────────
    def _save(self) -> tuple[str, str]:
        return self._save_doc(self.doc)

    def _save_doc(self, doc) -> tuple[str, str]:
        out_doc    = self.path.parent / f"{self.path.stem}_corregit.docx"
        out_report = self.path.parent / f"{self.path.stem}_informe.md"
        doc.save(str(out_doc))
        out_report.write_text(self.report.to_markdown(), encoding='utf-8')
        return str(out_doc), str(out_report)


# ─── Punt d'entrada ──────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description='InDret — Corrector de format d\'articles (v2)'
    )
    parser.add_argument('article',
        help='Fitxer .docx a corregir')
    parser.add_argument('--edicio', default='',
        metavar='NUM',
        help='Número d\'edició InDret, p. ex. "1/2024". S\'insereix a la plantilla.')
    parser.add_argument('--plantilla', default='',
        metavar='FITXER',
        help='Ruta a plantilla.docx. Si s\'omet, es cerca automàticament al directori del script.')
    parser.add_argument('--doi', default='',
        metavar='DOI',
        help='Identificador DOI de l\'article, p. ex. "10.31009/InDret.2024.i1.01".')
    parser.add_argument('--autor', default='',
        metavar='NOM',
        help='Nom de l\'autor/a (sobreescriu la detecció automàtica).')
    parser.add_argument('--recepcio', default='',
        metavar='DATA',
        help='Data de recepció de l\'article, p. ex. "12 de enero de 2025".')
    parser.add_argument('--acceptacio', default='',
        metavar='DATA',
        help='Data d\'acceptació de l\'article, p. ex. "3 de marzo de 2025".')
    parser.add_argument('--sense-plantilla', action='store_true',
        help='Força el mode clàssic (correcció directa sense plantilla).')
    args = parser.parse_args()

    if not os.path.exists(args.article):
        print(f"Error: fitxer no trobat → {args.article}")
        sys.exit(1)

    # Auto-detecció de la plantilla
    plantilla_path = args.plantilla
    if not plantilla_path and not args.sense_plantilla:
        candidate = Path(__file__).parent / 'plantilla.docx'
        if candidate.exists():
            plantilla_path = str(candidate)

    print(f"\n InDret — Corrector de format v2")
    print(f" Article:  {args.article}")
    if plantilla_path:
        print(f" Plantilla:   {plantilla_path}")
        print(f" Edició:      {args.edicio or '(no especificada)'}")
        print(f" DOI:         {args.doi or '(no especificat)'}")
        print(f" Autor:       {args.autor or '(detecció automàtica)'}")
        print(f" Recepció:    {args.recepcio or '(no especificada)'}")
        print(f" Acceptació:  {args.acceptacio or '(no especificada)'}")
    else:
        print(f" Mode:     correcció directa (sense plantilla)")
    print("─" * 50)

    corrector = InDretCorrector(args.article)

    if plantilla_path:
        out_doc, out_report = corrector.template_run(
            plantilla_path, args.edicio, args.doi,
            args.autor, args.recepcio, args.acceptacio
        )
    else:
        out_doc, out_report = corrector.run()

    print(f"\n[✓] Document generat → {out_doc}")
    print(f"[✓] Informe          → {out_report}\n")


if __name__ == '__main__':
    main()
